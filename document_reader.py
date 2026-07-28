import os
import re
import pdfplumber
from docx import Document

# أسماء الأقسام المعروفة داخل الـ CV
SECTION_NAMES = {
    "education": "Education",
    "experience": "Experience",
    "work experience": "Experience",
    "professional experience": "Experience",
    "skills": "Skills",
    "technical skills": "Skills",
    "projects": "Projects",
    "certifications": "Certifications",
    "certificates": "Certifications",
    "languages": "Languages",
    "summary": "Summary",
    "profile": "Profile",
    "objective": "Objective",
}

# رموز الـ Bullet
BULLET_SYMBOLS = "•▪◦●○■□◆►▶★✓✔"


def clean_text(text):
    """تنظيف النص المستخرج من الـ CV."""

    if not text:
        return ""

    for symbol in BULLET_SYMBOLS:
        text = text.replace(symbol, "")

    text = re.sub(r"\r", "", text)
    text = re.sub(r"\n{2,}", "\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)

    return text.strip()


def extract_pdf_text(file_path):
    """قراءة ملفات PDF."""

    text = ""

    try:
        with pdfplumber.open(file_path) as pdf:

            for page in pdf.pages:

                page_text = page.extract_text()

                if page_text:
                    text += page_text + "\n"

    except Exception as e:
        raise Exception(f"Error reading PDF: {e}")

    return clean_text(text)


def extract_word_text(file_path):
    """قراءة ملفات Word."""

    try:
        document = Document(file_path)

    except Exception as e:
        raise Exception(f"Error reading Word file: {e}")

    text = ""

    # الفقرات
    for paragraph in document.paragraphs:

        if paragraph.text.strip():
            text += paragraph.text.strip() + "\n"

    # الجداول
    for table in document.tables:

        for row in table.rows:

            values = []

            for cell in row.cells:

                if cell.text.strip():
                    values.append(cell.text.strip())

            if values:
                text += " | ".join(values) + "\n"

    return clean_text(text)


def extract_text(file_path):
    """
    استخراج النص من ملف PDF أو DOCX.
    """

    if not os.path.exists(file_path):
        raise FileNotFoundError(file_path)

    extension = os.path.splitext(file_path)[1].lower()

    if extension == ".pdf":
        return extract_pdf_text(file_path)

    elif extension == ".docx":
        return extract_word_text(file_path)

    else:
        raise Exception("Only PDF and DOCX files are supported.")


if __name__ == "__main__":

    path = input("Enter CV path: ")

    try:

        cv = extract_text(path)

        print("=" * 60)
        print(cv)
        print("=" * 60)

    except Exception as e:

        print(e)
print("document_reader loaded")
print(extract_text)